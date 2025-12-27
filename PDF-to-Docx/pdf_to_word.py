from pdf2docx import Converter
from docx import Document
from docx.shared import Pt, RGBColor
import os

pdf_folder = "PDF-to-Docx"
pdf_filename = "django_assignment.pdf"

pdf_path = os.path.join(pdf_folder, pdf_filename)
word_path = os.path.join(pdf_folder, "django_assignment.docx")

try:
    if not os.path.exists(pdf_path):
        print("PDF file not found!")
    else:
        cv = Converter(pdf_path)
        cv.convert(word_path)
        cv.close()
        
        doc = Document(word_path)
        
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.font.size:
                    run.font.size = Pt(run.font.size.pt + 1)
                else:
                    run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
                
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.font.size:
                                run.font.size = Pt(run.font.size.pt + 1)
                            else:
                                run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(0, 0, 0)
        
        doc.save(word_path)
        print("Conversion completed successfully!")
        
except Exception as e:
    print(f"Error: {e}")